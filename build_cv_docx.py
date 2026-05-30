"""Generate Sharjeel Asad Sheikh CV as Word document."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(11, 99, 217)
    return h


def add_bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    return p


def add_role(title, org_dates, bullets):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    p.add_run(f"\n{org_dates}").italic = True
    for b in bullets:
        add_bullet(b)
    doc.add_paragraph()


# Header
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Sharjeel Asad Sheikh")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(13, 27, 42)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run(
    "Aerospace Engineer · CFD Analyst · UAS & RC Specialist · Business Leader"
).font.size = Pt(11)

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.add_run(
    "Multan, Pakistan  |  0306-7320679  |  sharjeelasad871@gmail.com  |  "
    "linkedin.com/in/sharjeel-asad-sheikh-2412212aa"
).font.size = Pt(10)

doc.add_paragraph()

# Profile
add_heading("Profile", level=1)
doc.add_paragraph(
    "Aerospace Engineering undergraduate at NUST (2023–2027) with hands-on experience "
    "in CFD simulation, CAD-driven aerodynamics workflows, and building flight-tested "
    "UAS/RC platforms. Strong leadership across HR, operations, and business development "
    "in fast-moving student and startup teams."
)

# Experience
add_heading("Experience", level=1)

add_role(
    "FEA Engineer Intern",
    "SOCO Engineers GmbH · Germany",
    [
        "Performed finite element analysis on structural components; prepared meshes, boundary conditions, and load cases for simulation workflows.",
        "Evaluated stress, strain, and deformation results; supported design validation and technical reporting for engineering deliverables.",
    ],
)

add_role(
    "CFD Analysis Intern",
    "Computational Aeronautics Lab, Islamabad · May 2025 – Present",
    [
        "Simulated full aerodynamics of F1 Ferrari F10 with ANSYS Fluent & OpenFOAM; analyzed drag, lift, and pressure distributions; produced airflow visualizations and technical reports.",
        "Built CAD models, generated meshes, and cross-validated simulation results across both solvers.",
    ],
)

add_role(
    "Head of Business Development",
    "MACHX Educational Platform, Islamabad · Oct 2024 – Present",
    [
        "Directed a cross-functional team across scheduling, student support, marketing, and financial planning; scaled platform revenue through targeted acquisition strategies.",
        "Executed business strategy for offerings and engagement; handled budgeting, revenue tracking, and stakeholder reporting.",
    ],
)

# Projects
add_heading("Aerospace Projects & Competitions", level=1)

projects = [
    (
        "Engine Remaining Useful Life (RUL) Detector — AI",
        "Predictive Maintenance · Machine Learning",
        "Developed an AI-based remaining useful life estimator for aircraft engines using sensor and operational data. Cleaned time-series signals, engineered degradation features, and trained models to predict how much useful life remains before maintenance—helping teams move from reactive repairs to data-informed scheduling and lower unplanned downtime.",
    ),
    (
        "RC Plane Builder — DBFC 2025",
        "Mach X · GIKI (2025)",
        "Designed, fabricated, and flight-tested a competition-ready fixed-wing aircraft; optimized wing geometry, airframe structure, and propulsion.",
    ),
    (
        "Quadcopter Builder — UAS Challenge 2024",
        "Peregrine Technologies (2024)",
        "Built a fully functional quadcopter; handled frame design, flight-controller configuration, PID tuning, electronics integration, and field flight tests.",
    ),
    (
        "VTOL Aircraft Builder — UAS Challenge 2026",
        "Peregrine Technologies (2026)",
        "Engineered a fixed-wing/multirotor hybrid VTOL; led airframe fabrication, avionics integration, hybrid propulsion design, and transition-control testing.",
    ),
]

for name, meta, desc in projects:
    p = doc.add_paragraph()
    r = p.add_run(name)
    r.bold = True
    doc.add_paragraph(meta).runs[0].italic = True
    doc.add_paragraph(desc)
    doc.add_paragraph()

# Leadership
add_heading("Leadership & Roles", level=1)

add_role(
    "Head of Human Resources",
    "Peregrine Technologies · Present",
    [
        "Leading talent acquisition, team structuring, onboarding, performance evaluation frameworks, and culture development.",
    ],
)

add_role(
    "Vice President of Operations",
    "Aerothon — Drone Tech × NUST · Oct – Dec 2025",
    [
        "Managed end-to-end operations; coordinated multi-team logistics and external relations for seamless execution.",
    ],
)

p = doc.add_paragraph()
p.add_run("Societies & Event Roles").bold = True
doc.add_paragraph("NUST / Islamabad").runs[0].italic = True
for item in [
    "Executive External Relations — Piston Cup (SMME)",
    "Director Liaison — National Cultural Fest, Islamabad",
    "Orientation Guide — NUST Orientation 2024",
    "Executive Protocols — FICS (Innovative & Creative Solutions)",
    "Executive Protocols — National Literary Festival, Islamabad",
]:
    add_bullet(item)

doc.add_paragraph()

# Skills
add_heading("Skills", level=1)

skills_blocks = [
    ("Technical", "Aerodynamics · RC Aviation · UAS · Formula One CFD · Motorsports · Entrepreneurship\nANSYS Fluent · OpenFOAM"),
    ("Engineering Tools", "SolidWorks / CATIA · MATLAB & Numerical Methods · data visualization & reporting"),
    ("Management", "Team leadership & HR ops · strategic business planning · operations & resource management · financial planning & budgeting · digital marketing & outreach · stakeholder communication"),
]

for name, text in skills_blocks:
    p = doc.add_paragraph()
    p.add_run(name).bold = True
    doc.add_paragraph(text)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Languages").bold = True
for lang in [
    "Urdu — Native / C2",
    "English — Proficient / C2",
    "German — Upper Intermediate / B2",
]:
    add_bullet(lang)

doc.add_paragraph()

# Education
add_heading("Education", level=1)

add_role(
    "B.S. Aerospace Engineering",
    "NUST, Islamabad · 2023 – 2027",
    ["Focus areas: aerodynamics, propulsion, flight dynamics, and practical UAS engineering."],
)

p = doc.add_paragraph()
p.add_run("Intermediate FSc (Pre-Eng.)").bold = True
doc.add_paragraph("Kips College, Multan · 2021 – 2023").runs[0].italic = True
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run("Certifications").bold = True
for cert in [
    "ANSYS Fluent CFD",
    "OpenFOAM CFD Training",
    "SolidWorks / CATIA CAD",
    "MATLAB & Numerical Methods",
    "Project Management & Strategy",
    "Digital Marketing & Entrepreneurship",
]:
    add_bullet(cert)

out_path = r"d:\ai\Sharjeel_Asad_Sheikh_CV.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
